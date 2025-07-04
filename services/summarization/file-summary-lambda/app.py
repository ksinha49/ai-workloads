# ------------------------------------------------------------------------------
# app.py
# ------------------------------------------------------------------------------
"""
Module: app.py
Description:
  1. Receive pre-generated summaries from the state machine.
  2. Format summaries—including Markdown-style tables—into a Unicode-capable summary PDF.
  3. Fetch the original document from S3, merge summary pages before the original.
  4. Upload the merged document back to S3.
Pre- and post-conditions are documented on the main handler.


Version: 1.0.2
Created: 2025-05-05
Last Modified: 2025-06-28
Modified By: Koushik Sinha
"""

from __future__ import annotations

from common_utils import configure_logger, lambda_response
import re
import os
from io import BytesIO
from typing import Any, Dict, List, Union
from services.summarization.models import SummaryEvent
from datetime import datetime

import boto3
try:
    from botocore.exceptions import ClientError, BotoCoreError
except ModuleNotFoundError:  # pragma: no cover - fallback for minimal env
    class ClientError(Exception):
        pass

    class BotoCoreError(Exception):
        pass
from common_utils.get_ssm import (
    get_values_from_ssm,
    get_environment_prefix,
)
from fpdf import FPDF
from docx import Document
import json
import xml.etree.ElementTree as ET
from unidecode import unidecode

FONT_PATH = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")
BOLD_FONT_PATH = os.path.join(os.path.dirname(__file__), "DejaVuSans-Bold.ttf")

# Module Metadata
__author__ = "Koushik Sinha"
__version__ = "1.0.2"
__modified_by__ = "Koushik Sinha"


# ─── Logging Configuration ─────────────────────────────────────────────────────
logger = configure_logger(__name__)

_s3_client = boto3.client("s3")




def format_summary_content(raw: str) -> List[Union[str, List[List[str]]]]:
    """
    Parse raw summary text into a sequence of:
      - paragraph strings, or
      - tables as list-of-rows (each row is a list of cell strings).

    Supports Markdown-style tables. Example:

        | Col1 | Col2 |
        |------|------|
        | a    | b    |
        | c    | d    |

    Returns:
        A list of mixed paragraph/table blocks.
    """
    temp_str = raw.replace("'", "APOSTRO_PHE") 
    lines = temp_str.splitlines()
    blocks: List[Union[str, List[List[str]]]] = []
    i = 0
    while i < len(lines):
        if (
            lines[i].startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[-\s|]+\|?$", lines[i + 1])
        ):
            # Table start
            header = [c.strip() for c in lines[i].split("|")[1:-1]]
            i += 2
            rows: List[List[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].split("|")[1:-1]]
                if len(cells) != len(header):
                    logger.warning("Malformed table row: %s", lines[i])
                    break
                rows.append(cells)
                i += 1
            blocks.append([header] + rows)
        else:
            text = lines[i].strip()
            text = text.replace("APOSTRO_PHE", "'")
            if text:
                blocks.append(text)
            i += 1
    return blocks


def render_table(
    pdf: FPDF,
    table: List[List[str]],
    x: float,
    y: float,
    total_width: float,
) -> None:
    """
    Draw a table at (x, y) within the document.

    Args:
        pdf: FPDF instance.
        table: Rows (first row is header).
        x, y: Starting coordinates.
        total_width: Total horizontal space.

    Note on `multi_cell(..., ln=3)`:
      - ln=3 positions the next write immediately to the right
        of the last cell, keeping the same y-coordinate.
    """
    pdf.set_xy(x, y)
    cols = len(table[0])
    col_w = total_width / cols
    line_h = pdf.font_size*1.5
    # Header row
    prefix = get_environment_prefix()
    font_size = get_values_from_ssm(f"{prefix}/SUMMARY_PDF_FONT_SIZE")
    pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
    pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
    pdf.set_font("DejaVu", size=int(font_size))
    #pdf.set_font("Times", size=10)
    with pdf.table() as pdf_table:
      pdf_row = pdf_table.row()
      for cell in table[0]:
        pdf_row.cell(cell)
      for row in table[1:]:
        pdf.set_x(x)
        pdf_row = pdf_table.row()
        for cell in row:
            pdf_row.cell(cell)     
      pdf.set_margins(20, 10, 20)

def remove_asterisks(text):
  """Removes all occurrences of '*' and '**' from a string.

  Args:
    text: The input string.

  Returns:
    The string with '*' and '**' removed.
  """
  return re.sub(r'\*\*|\*', '', text)


def _add_title_page(pdf: FPDF, font_size: int, font_size_bold: int) -> None:
    """Create the initial title page."""
    pdf.add_page()
    pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
    pdf.set_font("DejaVu", style="B", size=int(font_size_bold))
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    pdf.multi_cell(w=150, h=5, text=current_date, align="R")
    pdf.ln(1)
    pdf.multi_cell(w=150, h=5, text="APS Summary", align="C")
    pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
    pdf.set_font("DejaVu", size=int(font_size))
    pdf.ln(2)


def _write_paragraph(pdf: FPDF, text: str, font_size: int, font_size_bold: int) -> None:
    """Write a single paragraph to the PDF respecting markdown markers."""
    text = unidecode(text)
    formatted = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    if text.startswith("**") and text.endswith("**"):
        pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
        pdf.set_font("DejaVu", style="B", size=int(font_size_bold))
        pdf.multi_cell(w=150, h=5, text=formatted)
    elif formatted.startswith("**") and not formatted.endswith("**"):
        formatted = remove_asterisks(formatted)
        pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
        pdf.set_font("DejaVu", size=int(font_size))
        pdf.multi_cell(w=150, h=5, text=formatted)
    elif not formatted.startswith("**") and formatted.endswith("**"):
        formatted = remove_asterisks(formatted)
        pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
        pdf.set_font("DejaVu", size=int(font_size))
        pdf.multi_cell(w=150, h=5, text=formatted)
    elif formatted.startswith("*"):
        formatted = remove_asterisks(formatted)
        pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
        pdf.set_font("DejaVu", size=int(font_size))
        pdf.multi_cell(w=150, h=5, text=formatted)
    else:
        if not formatted.startswith("Note:"):
            pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
            pdf.set_font("DejaVu", size=int(font_size))
            pdf.multi_cell(w=150, h=5, text=formatted)
    pdf.ln(2)


def _render_table(pdf: FPDF, table: List[List[str]]) -> None:
    """Render a markdown table."""
    render_table(pdf, table, x=20, y=pdf.get_y(), total_width=170)
    pdf.ln(2)

def create_summary_pdf(summaries: List[str]) -> BytesIO:
    """Build a summary PDF from a list of summary blocks."""

    prefix = get_environment_prefix()
    font_size = get_values_from_ssm(f"{prefix}/SUMMARY_PDF_FONT_SIZE")
    font_size_bold = get_values_from_ssm(f"{prefix}/SUMMARY_PDF_FONT_SIZE_BOLD")

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_margins(20, 20)
    buf = BytesIO()

    for idx, (title, raw) in enumerate(summaries):
        if idx == 0:
            _add_title_page(pdf, int(font_size), int(font_size_bold))
        else:
            if title != "NA":
                pdf.ln(5)
                pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
                pdf.set_font("DejaVu", style="B", size=int(font_size_bold))
                pdf.multi_cell(w=150, h=5, text=title, align="C")
                pdf.ln(3)
                pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
                pdf.set_font("DejaVu", size=int(font_size))

        blocks = format_summary_content(raw)
        for block in blocks:
            if isinstance(block, str):
                _write_paragraph(pdf, block, int(font_size), int(font_size_bold))
            else:
                _render_table(pdf, block)

    pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
    pdf.set_font("DejaVu", style="B", size=int(font_size_bold))
    pdf.ln(1)
    pdf.multi_cell(w=150, h=5, text="====End of APS Summary====", align="C")

    pdf.output(buf)
    buf.seek(0)
    return buf


def create_summary_docx(summaries: List[tuple[str, str]]) -> BytesIO:
    """Build a DOCX file containing the summaries."""

    doc = Document()
    for title, text in summaries:
        if title != "NA":
            doc.add_heading(title, level=1)
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_summary_json(summaries: List[tuple[str, str]]) -> BytesIO:
    """Serialize summaries to JSON."""

    data = [{"Title": t, "content": c} for t, c in summaries]
    buf = BytesIO(json.dumps(data).encode())
    buf.seek(0)
    return buf


def create_summary_xml(summaries: List[tuple[str, str]]) -> BytesIO:
    """Serialize summaries to XML."""

    root = ET.Element("summaries")
    for title, text in summaries:
        s = ET.SubElement(root, "summary")
        ET.SubElement(s, "title").text = title
        ET.SubElement(s, "content").text = text
    buf = BytesIO()
    tree = ET.ElementTree(root)
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    buf.seek(0)
    return buf

def upload_buffer_to_s3(buffer: BytesIO, bucket: str, bucket_key: str, content_type: str) -> None:
    """Upload a summary buffer to S3 with a specified content type."""

    try:
        _s3_client.put_object(
            Bucket=bucket,
            Key=bucket_key,
            Body=buffer.getvalue(),
            ContentType=content_type,
        )
    except (ClientError, BotoCoreError) as exc:
        logger.exception("Failed to upload summary file to S3")
        raise RuntimeError("Unable to upload summary") from exc
    logger.info("Uploaded summary to s3://%s/%s", bucket, bucket_key)


def process_for_summary(event: SummaryEvent, context: Any) -> Dict[str, Any]:
    """
    Main Lambda logic: summarize, merge, upload.

    Pre-conditions:
      event must include keys:
        - "collection_name": str
        - "statusCode": int
        - "organic_bucket": str
        - "organic_bucket_key": str

    Post-conditions on success:
      Returns the original event dict, plus:
        - "summary_bucket_name"
        - "summary_bucket_key"
        - "merged_bucket_name"
        - "merged_bucket_key"
        - "statusCode": 200
        - "statusMessage":...
    On failure:
      Returns {"statusCode":500, "statusMessage":<error>}.
    """
    event_body = event.to_dict()

    required = {"collection_name", "statusCode", "organic_bucket", "organic_bucket_key"}
    if not required.issubset(event_body):
        msg = f"Missing required event keys: {required}"
        logger.error(msg)
        return {"statusCode": 400, "statusMessage": msg}

    if event_body["statusCode"] != 200:
        upstream = event_body.get("statusMessage", "")
        msg = f"Upstream error: {upstream}"
        logger.error(msg)
        return {"statusCode": 500, "statusMessage": msg}

    try:
        if not isinstance(event_body.get("summaries"), list):
            raise RuntimeError("summaries list missing")

        summaries: List[tuple[str, str]] = []
        for item in event_body.get("summaries", []):
            title = item.get("Title", "")
            content = item.get("content", "")
            summaries.append((title, content))

        fmt = str(event_body.get("output_format", "pdf")).lower()
        if fmt == "pdf":
            summary_buf = create_summary_pdf(summaries)
            ext = "pdf"
            ctype = "application/pdf"
        elif fmt == "docx":
            summary_buf = create_summary_docx(summaries)
            ext = "docx"
            ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "json":
            summary_buf = create_summary_json(summaries)
            ext = "json"
            ctype = "application/json"
        elif fmt == "xml":
            summary_buf = create_summary_xml(summaries)
            ext = "xml"
            ctype = "application/xml"
        else:
            raise ValueError("Unsupported output format")

        organic_file_key = event_body['organic_bucket_key']
        organic_bucket_name = event_body['organic_bucket']
        summary_file_key = organic_file_key.replace('extracted', 'summary')
        summary_file_key = os.path.splitext(summary_file_key)[0] + f".{ext}"
        logger.info("organic_file_key:%s", organic_file_key)
        #new_folder_name = organic_file_folder[1]
        upload_buffer_to_s3(summary_buf, organic_bucket_name, summary_file_key, ctype)

        return {
            **event_body,
            "summary_bucket_name": organic_bucket_name,
            "summary_bucket_key": summary_file_key,
            "statusCode": 200,
            "statusMessage": "Summarization  PDF uploaded",
        }

    except FileNotFoundError as exc:
        logger.exception("Required resource missing")
        msg = f"Missing file: {exc}" if exc.filename else "Required file missing"
        return {"statusCode": 500, "statusMessage": msg}
    except (ClientError, BotoCoreError) as exc:
        logger.exception("AWS error during summary processing")
        return {"statusCode": 502, "statusMessage": "Storage service error"}
    except ValueError as exc:
        logger.exception("Invalid input for summary processing")
        return {"statusCode": 400, "statusMessage": str(exc)}
    except Exception as exc:
        logger.exception("Unexpected error during summary processing")
        return {"statusCode": 500, "statusMessage": "Internal server error"}




def lambda_handler(event: SummaryEvent | Dict[str, Any], context: Any) -> Dict[str, Any]:
    """Triggered by the state machine to produce a summary file.

    1. Formats the summary text into the requested output format and merges it with the source
       document from S3 when applicable.
    2. Uploads the merged file and returns its location.

    Returns a response dictionary with status and file path.
    """
    try:
        if isinstance(event, dict):
            try:
                event = SummaryEvent.from_dict(event)
            except ValueError as exc:
                logger.exception("Invalid request to lambda_handler")
                return lambda_response(400, {"statusMessage": str(exc)})
        body = process_for_summary(event, context)
        status = body.get("statusCode", 200)
        return lambda_response(status, body)
    except (ClientError, BotoCoreError) as exc:
        logger.exception("AWS error invoking Lambda")
        return lambda_response(502, {"statusMessage": "AWS service error"})
    except ValueError as exc:
        logger.exception("Invalid request to lambda_handler")
        return lambda_response(400, {"statusMessage": str(exc)})
    except Exception as exc:
        logger.exception("lambda_handler failed")
        return lambda_response(500, {"statusMessage": "Internal server error"})
